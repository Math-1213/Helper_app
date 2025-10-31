from flask import Flask, request, jsonify
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import json

app = Flask(__name__)

# Caminho absoluto da pasta onde o server.py está
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
SAVE_DIR = os.path.join(BASE_DIR, "listas")

# Cria a pasta "listas" se ainda não existir
os.makedirs(SAVE_DIR, exist_ok=True)

@app.route('/push', methods=['PUT'])
def receive_list():
    try:
        list_id = request.headers.get('id', 'sem_id')
        data = request.get_json(force=True)

        list_name = data.get("name", "Sem Nome")
        tasks = data.get("tasks", [])

        filename = os.path.join(SAVE_DIR, f"lista_{list_id}.docx")

        # Cria novo documento ou abre existente
        if os.path.exists(filename):
            doc = Document(filename)
        else:
            doc = Document()
            # título fixo no topo do documento
            header = doc.add_paragraph("Histórico de Atualizações")
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
            header.runs[0].bold = True
            header.runs[0].font.size = Pt(20)
            doc.add_paragraph()  # espaçamento

        # --- Nome da Lista ---
        title = doc.add_paragraph(list_name)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_title = title.runs[0]
        run_title.bold = True
        run_title.font.size = Pt(16)
        run_title.font.color.rgb = RGBColor(50, 150, 255)

        # --- Data e hora ---
        now = datetime.now().strftime("%d/%m/%Y %H:%M:%S")
        date_p = doc.add_paragraph(f"Atualizado em: {now}")
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_date = date_p.runs[0]
        run_date.italic = True
        run_date.font.size = Pt(10)
        run_date.font.color.rgb = RGBColor(130, 130, 130)

        doc.add_paragraph()  # espaçamento

        # --- Lista de tarefas ---
        for task in tasks:
            text = task.get("text", "").strip()
            checked = task.get("checked", False)
            if not text:
                continue

            p = doc.add_paragraph("• " + text)
            run = p.runs[0]
            run.font.size = Pt(12)
            if checked:
                run.font.strike = True
                run.font.color.rgb = RGBColor(150, 150, 150)
            else:
                run.font.color.rgb = RGBColor(30, 30, 30)

        # --- Separador visual (linha horizontal simulada) ---
        doc.add_paragraph()
        sep = doc.add_paragraph("─" * 50)
        sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sep.runs[0].font.size = Pt(9)
        sep.runs[0].font.color.rgb = RGBColor(180, 180, 180)
        doc.add_paragraph()

        # --- Salva o arquivo ---
        doc.save(filename)

        print(f"\n📄 Lista salva em: {filename}")
        print(json.dumps({"id": list_id, "name": list_name, "tasks": tasks}, indent=4, ensure_ascii=False))

        return jsonify({
            "status": "ok",
            "file": filename,
            "message": f"Lista {list_id} salva com sucesso"
        }), 200

    except Exception as e:
        print("❌ Erro ao processar requisição:", e)
        return jsonify({"status": "error", "message": str(e)}), 400


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)
